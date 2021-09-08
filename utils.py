import pandas as pd
from docx import Document
from docx.shared import Pt
from os import makedirs
from docx2pdf import convert


def get_sheet(excel_path):
    """
    returns the excel sheet class_name-marks.xlsx as Pandas dataframe 
    """
    sheet = pd.read_excel(excel_path)
    header = sheet.iloc[0]
    sheet = sheet[1:]
    sheet.columns = header

    cols_to_drop = ["sl no", "sl no.", "Sl No", "Sl No.", "Sl no", "Sl no.", "SL NO", "SL NO.", 
                    "adm no", "adm no.", "Adm no", "Adm no.", "Adm No", "Adm No.", "ADM NO", "ADM NO.", "ADMISSION NO", "ADMISSION NO."]
    sheet = sheet.drop(cols_to_drop, axis=1, errors="ignore")
    return sheet


def save_files(showGrades, convertToPdf, excel_path, save_dir, class_name, subjects, subject_details) :
    # Get excel sheet
    sheet = get_sheet(excel_path)
    print(sheet)

    # looping through each student
    for _, row in sheet.iterrows():
        full_name = row["Full name"]
        row = row.drop("Full name")
        print(full_name)
        print(row)
        print()
        marks = subject_details
        subjects_in_sheet = row.index

        for subject in subjects:
            for i in range(len(subjects_in_sheet)):
                if(subjects_in_sheet[i].upper()[:2] == "SS"):
                    subject_match = "Social Science" in subjects or "Ss" in subjects
                elif(subjects_in_sheet[i].upper()[:2] == "IT"):
                    subject_match = "Information Technology" in subjects or "It" in subjects
                else:
                    subject_match = subjects_in_sheet[i].lower().startswith(subject.lower()[:3])

                if subject_match:
                    mark = row[i]
                    mark_to_insert = " "; 
                    grade = " ";
                    if(str(mark) not in "----                     nan"):
                        mark = int(mark)
                        mark_to_insert = mark
                        if showGrades:
                            grade = get_grade(mark, marks[i][2])

                    marks[i][1], marks[i][3]= mark_to_insert, grade

        # adding document headings
        doc = Document()
        school = doc.add_heading("St. Thomas HSS, Nadavayal", 0)
        school.alignment = 1
        heading2 = doc.add_heading("First Terminal Evaluation", 0)
        heading2.alignment = 1
        class_heading = doc.add_heading(f"Class {class_name}", 1)
        class_heading.alignment = 1

        # adding name of student
        name_of_student = doc.add_paragraph().add_run(f"\n\nName of the student : {full_name}\n\n")
        name_of_student.font.name = "Arial"
        name_of_student.font.size = Pt(16)

        # creating table and adding column headings
        rows = 9
        cols = 4 if showGrades else 3
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "TableGrid"
        head_row = table.rows[0].cells
        head_row[0].text = "Subject"
        head_row[0].paragraphs[0].alignment = 1
        head_row[0].paragraphs[0].runs[0].font.size = Pt(15)
        head_row[1].text = "Marks received"
        head_row[1].paragraphs[0].alignment = 1
        head_row[1].paragraphs[0].runs[0].font.size = Pt(15)
        head_row[2].text = "Maximum marks"
        head_row[2].paragraphs[0].alignment = 1
        head_row[2].paragraphs[0].runs[0].font.size = Pt(15)
        if showGrades:
            head_row[3].text = "Grade"
            head_row[3].paragraphs[0].alignment = 1
            head_row[3].paragraphs[0].runs[0].font.size = Pt(15)

        # adding marks to table
        for sub, table_row in zip(marks, list(table.rows)[1:]):
            cells = table_row.cells
            for cell, detail in zip(cells, sub):
                cell.text = str(detail)
                cell.paragraphs[0].runs[0].font.size = Pt(14)
                cell.paragraphs[0].alignment = 1

        # # saving reports to respective folder
        makedirs(f"{save_dir}/{class_name}/", exist_ok=True)
        doc.save(f"{save_dir}/{class_name}/{full_name}.docx")

    # converting word document to pdf
    if convertToPdf:
        convert(f"{save_dir}/{class_name}/")


def get_grade(mark, max_mark):
    if(max_mark==80): 
        if(72<=mark<=80):
            grade = "A+"
        elif(64<=mark<=71):
            grade = "A"
        elif(56<=mark<=63):
            grade = "B+"
        elif(48<=mark<=55):
            grade = "B"
        elif(40<=mark<=47):
            grade = "C+"
        elif(32<=mark<=39):
            grade = "C"
        elif(24<=mark<=31):
            grade = "D+"
        elif(16<=mark<=23):
            grade = "D"
        else:
            grade = "E"
    elif(max_mark==40):
        if(36<=mark<=40):
            grade = "A+"
        elif(32<=mark<=35):
            grade = "A"
        elif(28<=mark<=31):
            grade = "B+"
        elif(24<=mark<=27):
            grade = "B"
        elif(20<=mark<=23):
            grade = "C+"
        elif(16<=mark<=19):
            grade = "C"
        elif(12<=mark<=15):
            grade = "D+"
        elif(8<=mark<=11):
            grade = "D"
        else:
            grade = "E"
    elif (max_mark==25):
        grade = "A+"
    elif(max_mark==20):
        if(18<=mark<=20):
            grade = "A+"
        elif(16<=mark<=17):
            grade = "A"
        elif(15<=mark<=16):
            grade = "B+"
        elif(13<=mark<=14):
            grade = "B"
        elif(11<=mark<=12):
            grade = "C+"
        elif(10<=mark<=11):
            grade = "C"
        elif(8<=mark<=9):
            grade = "D+"
        elif(6<=mark<=7):
            grade = "D"
        else:
            grade = "E"
    
    return grade
